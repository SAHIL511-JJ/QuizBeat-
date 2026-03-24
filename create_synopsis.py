from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

# Create a new Document
doc = Document()

# Set margins to 1 inch
sections = doc.sections
for section in sections:
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

def add_formatted_paragraph(text, size=12, bold=False, alignment=WD_ALIGN_PARAGRAPH.CENTER, spacing_before=0, spacing_after=0):
    """Helper function to add formatted paragraphs"""
    p = doc.add_paragraph()
    p.alignment = alignment
    run = p.add_run(text)
    font = run.font
    font.name = 'Times New Roman'
    font.size = Pt(size)
    font.bold = bold
    if spacing_before > 0:
        p.paragraph_format.space_before = Pt(spacing_before)
    if spacing_after > 0:
        p.paragraph_format.space_after = Pt(spacing_after)
    p.paragraph_format.line_spacing = 1.5
    return p

# Title Page
add_formatted_paragraph("Project Synopsis", 18, bold=True, spacing_after=6)
add_formatted_paragraph("On", 14, spacing_after=6)
add_formatted_paragraph("QuizBeat - AI-Powered Quiz Platform", 18, bold=True, spacing_after=12)

add_formatted_paragraph("Submitted to", 14, bold=True, spacing_before=12, spacing_after=6)
add_formatted_paragraph("Savitribai Phule Pune University", 14, bold=True, spacing_after=12)

add_formatted_paragraph("In Partial Fulfillment of the requirement of the award of the degree of", 12, spacing_before=12)
add_formatted_paragraph("Bachelor of Computer Application", 12)
add_formatted_paragraph("TYBCA - SCIENCE, Sem VI", 12)
add_formatted_paragraph("Academic Year 2024-25", 12, spacing_after=12)

add_formatted_paragraph("Submitted by", 14, bold=True, spacing_before=12, spacing_after=6)
add_formatted_paragraph("Mr. Sahil Lamture", 14, bold=True, spacing_after=12)

add_formatted_paragraph("Under the Guidance of", 14, bold=True, spacing_before=12, spacing_after=6)
add_formatted_paragraph("Dr. [Guide Name]", 14, spacing_after=12)

add_formatted_paragraph("Department of Computer Application", 12, spacing_before=12)
add_formatted_paragraph("Alandi (D), 412105", 12, spacing_after=24)

# Add page break
doc.add_page_break()

# 1. Title
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.LEFT
run = p.add_run("1. Title")
run.font.name = 'Times New Roman'
run.font.size = Pt(16)
run.font.bold = True
p.paragraph_format.line_spacing = 1.5

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
run = p.add_run("QuizBeat - AI-Powered Quiz Platform")
run.font.name = 'Times New Roman'
run.font.size = Pt(12)
p.paragraph_format.line_spacing = 1.5

# 2. Introduction
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.LEFT
run = p.add_run("2. Introduction")
run.font.name = 'Times New Roman'
run.font.size = Pt(16)
run.font.bold = True
p.paragraph_format.line_spacing = 1.5
p.paragraph_format.space_before = Pt(6)

intro_text = """Traditional quiz preparation methods require students to manually create questions from textbooks, which is time-consuming and often leads to incomplete coverage of study material. Students struggle with self-assessment, lack immediate feedback, and find studying monotonous without engagement mechanisms.

QuizBeat is a comprehensive, web-based AI-powered quiz platform designed to revolutionize the learning experience. Built using modern technologies including React, Vite, FastAPI, and Firebase, the system leverages Groq's Llama 70B AI model to automatically generate quizzes from uploaded textbooks. The platform provides real-time multiplayer quiz competitions similar to Kahoot, making learning interactive and engaging. The significance of this project lies in its practical approach to solving real-world educational challenges by combining artificial intelligence with gamification to enhance student engagement and learning outcomes."""

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
run = p.add_run(intro_text)
run.font.name = 'Times New Roman'
run.font.size = Pt(12)
p.paragraph_format.line_spacing = 1.5

# 3. Objectives
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.LEFT
run = p.add_run("3. Objectives")
run.font.name = 'Times New Roman'
run.font.size = Pt(16)
run.font.bold = True
p.paragraph_format.line_spacing = 1.5
p.paragraph_format.space_before = Pt(6)

objectives = [
    "To design and develop a functional web-based quiz platform with AI-powered question generation.",
    "To implement secure file upload functionality supporting PDF, Word, and text documents.",
    "To integrate Groq's Llama 70B model for intelligent quiz generation with customizable difficulty levels.",
    "To develop a real-time multiplayer quiz system with Kahoot-like gameplay mechanics.",
    "To build an interactive dashboard with quiz history, performance analytics, and leaderboards.",
    "To implement Google OAuth authentication for secure user management.",
    "To create a responsive user interface accessible across devices and browsers.",
    "To provide instant feedback with AI-generated explanations for correct and incorrect answers."
]

for obj in objectives:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.left_indent = Inches(0.25)
    run = p.add_run(f"• {obj}")
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)
    p.paragraph_format.line_spacing = 1.5

# 4. Scope
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.LEFT
run = p.add_run("4. Scope")
run.font.name = 'Times New Roman'
run.font.size = Pt(16)
run.font.bold = True
p.paragraph_format.line_spacing = 1.5
p.paragraph_format.space_before = Pt(6)

# 4.1 Included in Scope
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.LEFT
run = p.add_run("4.1 Included in Scope")
run.font.name = 'Times New Roman'
run.font.size = Pt(14)
run.font.bold = True
p.paragraph_format.line_spacing = 1.5

scope_included = [
    "AI Quiz Generation: Automatic question creation from uploaded textbooks with difficulty customization (Easy, Medium, Hard).",
    "Chapter Selection: Ability to narrow quiz scope to specific chapters or topics.",
    "Multiplayer Mode: Real-time Kahoot-style quiz competitions with team support.",
    "Results & Analytics: Comprehensive performance tracking with AI-generated explanations.",
    "User Authentication: Secure Google Sign-In integration.",
    "Quiz Management: Save, edit, and share custom quizzes.",
    "Responsive Design: Cross-platform compatibility for desktop and mobile devices."
]

for item in scope_included:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.left_indent = Inches(0.25)
    run = p.add_run(f"• {item}")
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)
    p.paragraph_format.line_spacing = 1.5

# 4.2 Excluded from Scope
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.LEFT
run = p.add_run("4.2 Excluded from Scope")
run.font.name = 'Times New Roman'
run.font.size = Pt(14)
run.font.bold = True
p.paragraph_format.line_spacing = 1.5

scope_excluded = [
    "Native mobile application development (iOS/Android apps).",
    "Offline mode functionality.",
    "Integration with Learning Management Systems (LMS) like Moodle or Canvas.",
    "Video or audio-based quiz content generation."
]

for item in scope_excluded:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.left_indent = Inches(0.25)
    run = p.add_run(f"• {item}")
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)
    p.paragraph_format.line_spacing = 1.5

# 5. Methodology
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.LEFT
run = p.add_run("5. Methodology")
run.font.name = 'Times New Roman'
run.font.size = Pt(16)
run.font.bold = True
p.paragraph_format.line_spacing = 1.5
p.paragraph_format.space_before = Pt(6)

methodology_text = "The project follows a modular development methodology combining Agile principles with a modern microservices architecture."

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
run = p.add_run(methodology_text)
run.font.name = 'Times New Roman'
run.font.size = Pt(12)
p.paragraph_format.line_spacing = 1.5

# 5.1 Architecture
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.LEFT
run = p.add_run("5.1 Architecture")
run.font.name = 'Times New Roman'
run.font.size = Pt(14)
run.font.bold = True
p.paragraph_format.line_spacing = 1.5

architecture = [
    "Frontend (Presentation Tier): React with Vite for fast, responsive user interface.",
    "Backend (Logic Tier): FastAPI (Python) for AI processing and API endpoints.",
    "Database (Data Tier): Firebase Firestore for user data and Firebase Realtime Database for live quiz sessions.",
    "AI Layer: Groq API with Llama 70B model for intelligent question generation."
]

for item in architecture:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.left_indent = Inches(0.25)
    run = p.add_run(f"• {item}")
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)
    p.paragraph_format.line_spacing = 1.5

# 5.2 Tools and Technologies
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.LEFT
run = p.add_run("5.2 Tools and Technologies")
run.font.name = 'Times New Roman'
run.font.size = Pt(14)
run.font.bold = True
p.paragraph_format.line_spacing = 1.5

# Add table
table = doc.add_table(rows=7, cols=2)
table.style = 'Table Grid'

# Header row
header_cells = table.rows[0].cells
header_cells[0].text = 'Component'
header_cells[1].text = 'Technology'

# Data rows
data = [
    ('Frontend Framework', 'React + Vite'),
    ('Backend Framework', 'FastAPI (Python 3.9+)'),
    ('Database', 'Firebase (Firestore + Realtime Database)'),
    ('Authentication', 'Firebase Auth (Google OAuth)'),
    ('AI Model', 'Groq API (Llama 70B)'),
    ('Deployment', 'Vercel (Frontend) + Render (Backend)')
]

for i, (component, tech) in enumerate(data, start=1):
    row_cells = table.rows[i].cells
    row_cells[0].text = component
    row_cells[1].text = tech

# Format table text
for row in table.rows:
    for cell in row.cells:
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.name = 'Times New Roman'
                run.font.size = Pt(12)
            paragraph.paragraph_format.line_spacing = 1.5

# 6. Literature Review
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.LEFT
run = p.add_run("6. Literature Review")
run.font.name = 'Times New Roman'
run.font.size = Pt(16)
run.font.bold = True
p.paragraph_format.line_spacing = 1.5
p.paragraph_format.space_before = Pt(6)

lit_review = """Educational technology has evolved from traditional classroom methods to sophisticated AI-powered learning platforms. Research shows that gamification significantly improves student engagement and knowledge retention (Deterding et al., 2011). Modern quiz platforms like Kahoot have demonstrated the effectiveness of competitive learning environments. Recent advances in Large Language Models (LLMs) have enabled automated content generation with high accuracy. This project builds upon these established frameworks by combining AI-powered content generation with real-time multiplayer engagement, creating a cost-effective solution tailored for student self-assessment and collaborative learning."""

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
run = p.add_run(lit_review)
run.font.name = 'Times New Roman'
run.font.size = Pt(12)
p.paragraph_format.line_spacing = 1.5

# 7. Implementation Plan
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.LEFT
run = p.add_run("7. Implementation Plan")
run.font.name = 'Times New Roman'
run.font.size = Pt(16)
run.font.bold = True
p.paragraph_format.line_spacing = 1.5
p.paragraph_format.space_before = Pt(6)

# Add implementation table
impl_table = doc.add_table(rows=6, cols=3)
impl_table.style = 'Table Grid'

# Header row
header_cells = impl_table.rows[0].cells
header_cells[0].text = 'Phase'
header_cells[1].text = 'Tasks'
header_cells[2].text = 'Duration'

# Data rows
impl_data = [
    ('Phase 1: Planning', 'Requirement analysis, database schema design, and API planning', 'Week 1'),
    ('Phase 2: Core Dev', 'User authentication, file upload, and AI integration setup', 'Week 2-3'),
    ('Phase 3: Features', 'Quiz generation, results display, and explanation system', 'Week 4-5'),
    ('Phase 4: Multiplayer', 'Real-time game mechanics, leaderboards, and team support', 'Week 6-7'),
    ('Phase 5: Testing', 'Bug fixing, performance optimization, and deployment', 'Week 8')
]

for i, (phase, tasks, duration) in enumerate(impl_data, start=1):
    row_cells = impl_table.rows[i].cells
    row_cells[0].text = phase
    row_cells[1].text = tasks
    row_cells[2].text = duration

# Format table text
for row in impl_table.rows:
    for cell in row.cells:
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.name = 'Times New Roman'
                run.font.size = Pt(12)
            paragraph.paragraph_format.line_spacing = 1.5

# 8. Expected Outcomes
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.LEFT
run = p.add_run("8. Expected Outcomes")
run.font.name = 'Times New Roman'
run.font.size = Pt(16)
run.font.bold = True
p.paragraph_format.line_spacing = 1.5
p.paragraph_format.space_before = Pt(6)

outcomes = [
    "A fully functional web-based platform accessible via modern browsers.",
    "AI-powered quiz generation with 90%+ accuracy in question relevance.",
    "Real-time multiplayer quiz system supporting multiple concurrent games.",
    "Comprehensive analytics dashboard with performance tracking and insights.",
    "Secure authentication system with Google OAuth integration.",
    "Responsive design optimized for desktop and mobile devices."
]

for outcome in outcomes:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.left_indent = Inches(0.25)
    run = p.add_run(f"• {outcome}")
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)
    p.paragraph_format.line_spacing = 1.5

# 9. Conclusion
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.LEFT
run = p.add_run("9. Conclusion")
run.font.name = 'Times New Roman'
run.font.size = Pt(16)
run.font.bold = True
p.paragraph_format.line_spacing = 1.5
p.paragraph_format.space_before = Pt(6)

conclusion = """QuizBeat addresses the need for an intelligent, engaging, and accessible quiz platform for students. By leveraging cutting-edge AI technology with Groq's Llama 70B model and combining it with real-time multiplayer features, the project demonstrates the practical application of modern web development and artificial intelligence in solving educational challenges. The system covers the entire learning cycle from content upload to performance analysis, making it a valuable tool for students seeking efficient and engaging study methods."""

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
run = p.add_run(conclusion)
run.font.name = 'Times New Roman'
run.font.size = Pt(12)
p.paragraph_format.line_spacing = 1.5

# 10. References
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.LEFT
run = p.add_run("10. References")
run.font.name = 'Times New Roman'
run.font.size = Pt(16)
run.font.bold = True
p.paragraph_format.line_spacing = 1.5
p.paragraph_format.space_before = Pt(6)

references = [
    'Deterding, S., Dixon, D., Khaled, R., & Nacke, L., "From game design elements to gamefulness: defining gamification", Proceedings of the 15th International Academic MindTrek Conference, 2011.',
    'Goodfellow, I., Bengio, Y., & Courville, A., "Deep Learning", MIT Press, 2016.',
    'Meta AI, "React Documentation", https://react.dev, 2024.',
    'Groq, "Groq API Documentation", https://groq.com/docs, 2024.',
    'Firebase, "Firebase Documentation", https://firebase.google.com/docs, 2024.'
]

for i, ref in enumerate(references, start=1):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    run = p.add_run(f"{i}. {ref}")
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)
    p.paragraph_format.line_spacing = 1.5

# Save the document
doc.save('QuizBeat_Project_Synopsis_v2.docx')
print("Synopsis document created successfully!")
